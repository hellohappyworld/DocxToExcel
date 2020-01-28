# coding:utf-8
import os
from docx import Document
from openpyxl import load_workbook
import sys

reload(sys)
sys.setdefaultencoding('utf-8')


def readExcle():
    workbook = load_workbook(dirPath)
    sheet1 = workbook.get_sheet_by_name(writeSheet)
    print sheet1.max_row


def writeExcle():
    workbook = load_workbook(excelPath)
    active = workbook.get_sheet_by_name(writeSheet)
    maxRow = active.max_row
    active.cell(maxRow + 1, 1, kv["订单编号"])
    active.cell(maxRow + 1, 2, kv["供货单号"])
    active.cell(maxRow + 1, 3, kv["项目单位/收货方"])
    active.cell(maxRow + 1, 4, kv["项目名称"])
    active.cell(maxRow + 1, 5, kv["货物名称"])
    active.cell(maxRow + 1, 6, kv["数量"])
    active.cell(maxRow + 1, 7, kv["分项价（含税）"])
    active.cell(maxRow + 1, 8, kv["交货时间"])
    active.cell(maxRow + 1, 9, kv["支付比例"])
    workbook.save(excelPath)


def getText():
    file = Document(filePath)
    for para in file.paragraphs:
        if "本供货单合同货物价款支付比例按照" in para.text:
            kv["支付比例"] = para.text[18:-6]


def getTable():
    file = Document(filePath)
    tables = file.tables
    orderTable = tables[0]
    projectTable = tables[1]
    for row in orderTable.rows:
        for cell in row.cells:
            text = cell.text
            if "订单编号" in text:
                order = text[5:]
                kv["订单编号"] = order
            elif "供货单号" in text:
                supply = text[5:]
                kv["供货单号"] = supply
    for row in projectTable.rows[1:-1]:
        cells = row.cells
        kv["项目单位/收货方"] = cells[1].text
        kv["项目名称"] = cells[2].text
        kv["货物名称"] = cells[3].text
        kv["数量"] = cells[5].text
        kv["分项价（含税）"] = cells[9].text
        kv["交货时间"] = cells[10].text
        writeExcle()


if __name__ == "__main__":
    dirPath = sys.argv[1].decode("gbk")
    excelPath = sys.argv[2].decode("gbk")
    writeSheet = sys.argv[3].decode("gbk")
    print "docx路径-->".decode("utf-8"), dirPath
    print "excel路径-->".decode("utf-8"), excelPath
    print "excel Sheet-->".decode("utf-8"), writeSheet
    kv = {}
    for filesTup in os.walk(dirPath):
        root = filesTup[0]
        files = filesTup[2]
        for file in files:
            filePath = root + "\\" + file
            try:
                getText()
                getTable()
                print filePath, "succeed!"
            except Exception, e:
                print filePath, "fail!", str(e)
                continue
